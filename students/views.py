from datetime import date, timedelta
from io import BytesIO

import openpyxl
from django.contrib.auth.decorators import login_required, user_passes_test
from django.db.models import Q
from django.http import FileResponse
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.shortcuts import render
from django.utils.decorators import method_decorator
from django.views.generic import ListView
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from students.forms import (StudentForm, AntisocialBehaviorForm, IncentivesForm,
                            IndividualWorkForm, WorkWithParentsForm, SpecialistRecomendationsForm)
from students.models import Student, Group


def is_curator(user):
    return user.is_authenticated and user.role == 'Куратор'

def is_social_teacher(user):
    return user.is_authenticated and user.role == 'Социальный педагог'

def is_psychologist(user):
    return user.is_authenticated and user.role == 'Педагог-психолог'


@method_decorator(login_required, name='dispatch')
@method_decorator(user_passes_test(is_curator), name='dispatch')
class StudentListView(ListView):
    model = Student
    template_name = 'pages/main_curator_page.html'
    context_object_name = 'students'
    ordering = ['last_name']

    def get_queryset(self):
        return super().get_queryset().filter(group_number=self.request.user.group_number)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['group'] = self.request.user.group_number
        return context


@method_decorator(login_required, name='dispatch')
@method_decorator(user_passes_test(is_social_teacher), name='dispatch')
class SocialTeacherListView(ListView):
    model = Group
    template_name = 'pages/social_teacher.html'
    context_object_name = 'groups'
    ordering = 'group_number'


@method_decorator(login_required, name='dispatch')
@method_decorator(user_passes_test(is_psychologist), name='dispatch')
class PsychologistListView(ListView):
    model = Group
    template_name = 'pages/psychologist.html'
    context_object_name = 'groups'
    ordering = 'group_number'


def StudentDetailView(request, last_name, group):
    if request.method == 'POST':
        if 'form1-submit' in request.POST:
            form1 = StudentForm(instance=get_object_or_404(Student, last_name=last_name), data=request.POST, files=request.FILES)
            if form1.is_valid():
                student = get_object_or_404(Student, last_name=last_name)
                if bool(request.FILES) and student.image:
                    student.image.delete()
                form1.save()
        elif 'form2-submit' in request.POST:
            form2 = IncentivesForm(request.POST)
            if form2.is_valid():
                student = get_object_or_404(Student, last_name=last_name)
                incentives_instance = form2.save(commit=False)
                incentives_instance.save()
                student.incentives.add(incentives_instance)
        elif 'form3-submit' in request.POST:
            form3 = AntisocialBehaviorForm(request.POST)
            if form3.is_valid():
                student = get_object_or_404(Student, last_name=last_name)
                antisocial_instance = form3.save(commit=False)
                antisocial_instance.save()
                student.antisocial_behavior.add(antisocial_instance)
        elif 'form4-submit' in request.POST:
            form4 = SpecialistRecomendationsForm(request.POST)
            if form4.is_valid():
                student = get_object_or_404(Student, last_name=last_name)
                specialist_recommendations_instance = form4.save(commit=False)
                specialist_recommendations_instance.save()
                student.specialist_recommendations.add(specialist_recommendations_instance)
        elif 'form5-submit' in request.POST:
            form5 = IndividualWorkForm(request.POST)
            if form5.is_valid():
                student = get_object_or_404(Student, last_name=last_name)
                individualwork_instance = form5.save(commit=False)
                individualwork_instance.save()
                student.individualwork.add(individualwork_instance)
        elif 'form6-submit' in request.POST:
            form6 = WorkWithParentsForm(request.POST)
            if form6.is_valid():
                student = get_object_or_404(Student, last_name=last_name)
                workwithparents_instance = form6.save(commit=False)
                workwithparents_instance.save()
                student.workwithparents.add(workwithparents_instance)
    else:
        form = StudentForm(instance=get_object_or_404(Student, last_name=last_name))
        antisocialbehaviorform = AntisocialBehaviorForm()
        incentivesform = IncentivesForm()
        specialistrecomendationform = SpecialistRecomendationsForm()
        individualworkform = IndividualWorkForm()
        workwithparentsform = WorkWithParentsForm()

    context = {
        'form': StudentForm(instance=get_object_or_404(Student, last_name=last_name)),
        'antisocialbehaviorform': AntisocialBehaviorForm(),
        'incentivesform': IncentivesForm(),
        'specialistrecomendationform': SpecialistRecomendationsForm(),
        'individualworkform': IndividualWorkForm(),
        'workwithparentsform': WorkWithParentsForm(),
        'student': get_object_or_404(Student, last_name=last_name),
        'incentives': get_object_or_404(Student, last_name=last_name).incentives.all(),
        'antisocialbehaviors': get_object_or_404(Student, last_name=last_name).antisocial_behavior.all(),
        'specialistrecommendations': get_object_or_404(Student, last_name=last_name).specialist_recommendations.all(),
        'individualworks': get_object_or_404(Student, last_name=last_name).individualwork.all(),
        'workwithparents': get_object_or_404(Student, last_name=last_name).workwithparents.all(),
        'group': get_object_or_404(Student, last_name=last_name).group_number
    }
    return render(request, 'pages/student_page.html', context)



class AllStudentsFromGroupListView(ListView):
    template_name = 'pages/students_list_page.html'
    context_object_name = 'students'

    def get_queryset(self):
        group = self.kwargs['group']
        return Student.objects.filter(group_number=group).order_by('last_name')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        group = self.kwargs['group']
        context['group'] = group
        return context


@login_required
def show_group_report(request, group):
    date_today = date.today()
    eighteen_year_ago = date_today - timedelta(days=18*365)
    legal_age = 18
    legal_year = date_today.year - legal_age
    date_dict = dict(day=str(date_today.day).zfill(2),
                     month=str(date_today.month).zfill(2),
                     year=date_today.year
                     )
    minor_students_count = Student.objects.filter(
        Q(dateBirth__year__gt=legal_year) |
        Q(dateBirth__year=legal_year, dateBirth__month__gt=date_today.month) |
        Q(dateBirth__year=legal_year, dateBirth__month=date_today.month, dateBirth__day__gt=date_today.day),
        Q(group_number=group))
    table_data = (
        Student.objects.filter(group_number=group).count(),
        Student.objects.filter(Q(education_type='BUDGET'), Q(group_number=group)).count(),
        Student.objects.filter(Q(education_type='UNBUDGET'), Q(group_number=group)).count(),
        Student.objects.filter(Q(sex='MALE'), Q(group_number=group)).count(),
        Student.objects.filter(Q(sex='FEMALE'), Q(group_number=group)).count(),
        minor_students_count.count(),
        Student.objects.filter(Q(dateBirth__lte=eighteen_year_ago), Q(group_number=group)).count(),
        Student.objects.filter(Q(group_number=group)).exclude(Q(place_living__contains='Минск') |
                                                              Q(place_living='') |
                                                              Q(place_living__isnull=True)).count(),
        Student.objects.filter(Q(hostel='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(group_number=group)).exclude(Q(citizenship='Республика Беларусь')).count(),
        Student.objects.filter(Q(type_of_family='FULL'), Q(group_number=group)).count(),
        Student.objects.filter(Q(family_large='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(guardianship='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(family_foster='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(low_income_family='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(refugees='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(settlers='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(family_students='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(have_children_students='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(upo_students='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(ngz_students='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(sop='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(ipr='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(orphan_students='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(lica_iz_chisla_detei='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(deti='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(lica_iz_detei='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(brsm='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(prof_souz='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(self_government_student='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(invalid='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(oprf='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(circle='YES'), Q(group_number=group)).count(),
        Student.objects.filter(Q(circle='NO'), Q(group_number=group)).count(),
    )

    if request.method == 'POST':
        if 'form1-submit' in request.POST:
            doc = Document()
            head = doc.add_paragraph(f'СОЦИАЛЬНО-ПЕДАГОГИЧЕСКАЯ ХАРАКТЕРИСТИКА группы № {group}')
            mgkct = doc.add_paragraph(f'УО "МГК цифровых технологий" на {date_dict["day"]}.{date_dict["month"]}.{date_dict["year"]}г.')
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            head.style.font.size = Pt(14)
            head.style.font.name = "Times New Roman"
            mgkct.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mgkct.style.font.size = Pt(14)
            mgkct.style.font.name = "Times New Roman"
            table = doc.add_table(rows=3, cols=35)
            table.style = 'Table Grid'
            row2 = table.rows[1].cells
            row2[0].text = f'На\n' \
                           f'{date_dict["day"]}.' \
                           f'{date_dict["month"]}.'
            row2[1].text = 'Общее'
            row2[2].text = 'Бюджет'
            row2[3].text = 'Внебюджет'
            row2[4].text = 'Юноши'
            row2[5].text = 'Девушки'
            row2[6].text = 'Несовершеннолетние'
            row2[7].text = 'Совершеннолетние'
            row2[8].text = 'Иногородние'
            row2[9].text = 'Проживают в общежитии'
            row2[10].text = 'Иностранные граждане, вид на жительство'
            row2[11].text = 'Неполная'
            row2[12].text = 'Многодетная'
            row2[13].text = 'Попечительство'
            row2[14].text = 'Приемная'
            row2[15].text = 'Малообеспеченная'
            row2[16].text = 'Беженцы, дополнительная защита'
            row2[17].text = 'Переселенцы(чернобыльцы)'
            row2[18].text = 'Семейные учащиеся'
            row2[19].text = 'Имеют детей'
            row2[20].text = 'Находятся на государственном обеспечении в УПО'
            row2[21].text = 'НГЗ(отобраны у родителей но решения суда нет)'
            row2[22].text = 'СОП'
            row2[23].text = 'ИПР'
            row2[24].text = 'Дети-сироты'
            row2[25].text = 'Лица из числа детей-сирот'
            row2[26].text = 'Дети, оставшиеся без попечения родителей'
            row2[27].text = 'Лица из числа детей, оставшихся без попечения родителей'
            row2[28].text = 'Члены БРСМ'
            row2[29].text = 'Члены профсоюза'
            row2[30].text = 'Члены актива ученического самоуправления(актив)'
            row2[31].text = 'Инвалиды'
            row2[32].text = 'ОПФР'
            row2[33].text = 'Да'
            row2[34].text = 'Нет'
            row2 = table.rows[0].cells
            row2[0].text = ''
            row2[1].text = 'Кол-во учащихся'
            row2[12].text = 'Характеристика семей'
            row2[20].text = 'Характеристика учащихся'
            row2[28].text = 'Общественная занятость'
            row2[31].text = 'Группа здоровья'
            row2[33].text = 'Занятость в кружках'
            row3 = table.rows[2].cells

            for i in range(0, 34):
                cell = row3[i+1]
                cell.text = str(table_data[i])

            row2[0].merge(row2[0])
            row2[1].merge(row2[10])
            row2[11].merge(row2[19])
            row2[20].merge(row2[27])
            row2[28].merge(row2[30])
            row2[31].merge(row2[32])
            row2[33].merge(row2[34])

            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_section(start_type=WD_SECTION.NEW_PAGE)
            doc.sections[0].page_width = doc.sections[1].page_height
            doc.sections[0].page_height = doc.sections[1].page_width
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            response = FileResponse(file_stream, as_attachment=True, filename=f'Отчет по группе №{group}.docx')
            return response
        if 'form2-submit' in request.POST:
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet['A1'] = 'Общее'
            sheet['B1'] = 'Бюджет'
            sheet['C1'] = 'Внебюджет'
            sheet['D1'] = 'Юноши'
            sheet['E1'] = 'Девушки'
            sheet['F1'] = 'Несовершеннолетние'
            sheet['G1'] = 'Совершеннолетние'
            sheet['H1'] = 'Иногородние'
            sheet['I1'] = 'Проживают в общежитии'
            sheet['J1'] = 'Иностранные граждане, вид на жительство'
            sheet['K1'] = 'Неполная'
            sheet['L1'] = 'Многодетная'
            sheet['M1'] = 'Попечительство'
            sheet['N1'] = 'Приемная'
            sheet['O1'] = 'Малообеспеченная'
            sheet['P1'] = 'Беженцы, дополнительная защита'
            sheet['Q1'] = 'Переселенцы(чернобыльцы)'
            sheet['R1'] = 'Семейные учащиеся'
            sheet['S1'] = 'Имеют детей'
            sheet['T1'] = 'Находятся на государственном обеспечении в УПО'
            sheet['U1'] = 'НГЗ(отобраны у родителей но решения суда нет)'
            sheet['V1'] = 'СОП'
            sheet['W1'] = 'ИПР'
            sheet['X1'] = 'Дети-сироты'
            sheet['Y1'] = 'Лица из числа детей-сирот'
            sheet['Z1'] = 'Дети, оставшиеся без попечения родителей'
            sheet['AA1'] = 'Лица из числа детей, оставшихся без попечения родителей'
            sheet['AB1'] = 'Члены БРСМ'
            sheet['AC1'] = 'Члены профсоюза'
            sheet['AD1'] = 'Члены актива ученического самоуправления(актив)'
            sheet['AE1'] = 'Инвалиды'
            sheet['AF1'] = 'ОПФР'
            sheet['AG1'] = 'Да'
            sheet['AH1'] = 'Нет'
            sheet['A2'] = table_data[0]
            sheet['B2'] = table_data[1]
            sheet['C2'] = table_data[2]
            sheet['D2'] = table_data[3]
            sheet['E2'] = table_data[4]
            sheet['F2'] = table_data[5]
            sheet['G2'] = table_data[6]
            sheet['H2'] = table_data[7]
            sheet['I2'] = table_data[8]
            sheet['J2'] = table_data[9]
            sheet['K2'] = table_data[10]
            sheet['L2'] = table_data[11]
            sheet['M2'] = table_data[12]
            sheet['N2'] = table_data[13]
            sheet['O2'] = table_data[14]
            sheet['P2'] = table_data[15]
            sheet['Q2'] = table_data[16]
            sheet['R2'] = table_data[17]
            sheet['S2'] = table_data[18]
            sheet['T2'] = table_data[19]
            sheet['U2'] = table_data[20]
            sheet['V2'] = table_data[21]
            sheet['W2'] = table_data[22]
            sheet['X2'] = table_data[23]
            sheet['Y2'] = table_data[24]
            sheet['Z2'] = table_data[25]
            sheet['AA2'] = table_data[26]
            sheet['AB2'] = table_data[27]
            sheet['AC2'] = table_data[28]
            sheet['AD2'] = table_data[29]
            sheet['AE2'] = table_data[30]
            sheet['AF2'] = table_data[31]
            sheet['AG2'] = table_data[32]
            sheet['AH2'] = table_data[33]

            for col in sheet.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                sheet.column_dimensions[column].width = adjusted_width

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = f'attachment; filename=Отчет по группе №{group}.xlsx'

            workbook.save(response)

            return response

    context = {
        'group': group,
        'date': date_dict,
        'report_data': table_data
    }
    return render(request, 'pages/group_report_page.html', context=context)
